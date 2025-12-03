import os
import json
import argparse
from datetime import datetime

def load_json(path):
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def fmt(v):
    try:
        if isinstance(v, (int, float)):
            return f"{v:.4f}"
        return str(v)
    except Exception:
        return str(v)

def build_markdown(data_dir: str) -> str:
    artifacts = os.path.join(data_dir, 'artifacts')
    eval_path = os.path.join(artifacts, 'evaluation_summary.json')
    gnn_metrics = os.path.join(artifacts, 'gnn_metrics.json')
    mlp_metrics = os.path.join(artifacts, 'mlp_metrics.json')
    figures_dir = os.path.join(artifacts, 'figures')

    summary = load_json(eval_path) or {}
    gnn = load_json(gnn_metrics) or {}
    mlp = load_json(mlp_metrics) or {}

    mlp_variant_files = []
    if os.path.isdir(artifacts):
        for f in os.listdir(artifacts):
            if f.startswith('mlp_metrics_') and f.endswith('.json'):
                mlp_variant_files.append(os.path.join(artifacts, f))
    mlp_variants = []
    for p in sorted(mlp_variant_files):
        m = load_json(p) or {}
        name = os.path.basename(p).replace('mlp_metrics_', '').replace('.json', '')
        mlp_variants.append({
            'name': name,
            'rmse': m.get('rmse'),
            'mae': m.get('mae'),
            'mape': m.get('mape'),
            'train_time_sec': m.get('train_time_sec'),
            'epochs': m.get('epochs'),
            'hidden_dim': m.get('hidden_dim'),
            'lr': m.get('lr'),
            'features': m.get('features'),
        })

    gnn_inf_time = None
    if 'gnn' in summary and isinstance(summary['gnn'], dict):
        gnn_inf_time = summary['gnn'].get('inference_time_sec')
    gnn_hidden_dim = None
    gnn_num_layers = None
    ckpt_path = os.path.join(artifacts, 'gnn_model.pt')
    if os.path.exists(ckpt_path):
        try:
            import torch
            sd = torch.load(ckpt_path, map_location='cpu')
            if isinstance(sd, dict):
                try:
                    gnn_hidden_dim = sd['node_encoder.0.weight'].shape[0]
                except Exception:
                    gnn_hidden_dim = None
                try:
                    idxs = set()
                    for k in sd.keys():
                        if k.startswith('convs.'):
                            i = int(k.split('.')[1])
                            idxs.add(i)
                    gnn_num_layers = (max(idxs) + 1) if idxs else None
                except Exception:
                    gnn_num_layers = None
        except Exception:
            pass

    # image relative paths from artifacts/ to figures/
    fig_inference = os.path.join('figures', 'inference_times.png')
    fig_mlp_metrics = os.path.join('figures', 'mlp_metrics.png')
    fig_mlp_scatter = os.path.join('figures', 'mlp_scatter_coords_diff.png')
    fig_mlp_error = os.path.join('figures', 'mlp_error_coords_diff.png')
    fig_error_bins = os.path.join('figures', 'error_bins.png')
    fig_expansion = os.path.join('figures', 'expansion_bars.png')

    lines = []
    lines.append(f"# 基于图神经网络的最短路径距离近似研究")
    lines.append("")
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # 摘要
    lines.append("## 摘要")
    lines.append("本文在真实道路图（旧金山）上比较 Dijkstra、A* 与 GNN（GraphSAGE）在最短路径距离预测上的性能与效率。结果显示：Dijkstra 与 A* 在精度上近乎无误差，但推理耗时显著；GNN 推理速度接近毫秒级，适合实时近似场景，但当前配置下误差较大，仍需进一步调参优化。")

    # 引言
    lines.append("## 引言")
    lines.append("最短路径计算在导航、物流与交通仿真中至关重要。随着图规模增长，传统精确算法在实时场景的耗时成为瓶颈。本文探索以 GNN 学习近似距离，实现速度与精度的权衡。")

    # 方法
    lines.append("## 方法")
    lines.append("- 基线：Dijkstra（精确）、A*（地理启发式）。")
    lines.append("- GNN：GraphSAGE K 层编码节点嵌入，拼接源/目标嵌入经 MLP 回归距离。")
    lines.append("- 特征工程：节点经纬度、度数、PageRank、是否路口；Landmark 距离特征。")
    lines.append("- 评估：RMSE/MAE/MAPE 与推理耗时。")

    # 实验
    lines.append("## 实验与结果")
    lines.append(f"数据目录：`{data_dir}`；综合评估：`artifacts/evaluation_summary.json`。")
    lines.append("")
    # 指标表格（Markdown）
    lines.append("### 指标总览")
    lines.append("算法 | RMSE | MAE | MAPE | 推理耗时(s)")
    lines.append("--- | ---: | ---: | ---: | ---:")
    for key in ['dijkstra', 'a_star', 'mlp', 'gnn']:
        if key in summary:
            v = summary[key]
            lines.append(f"{key} | {fmt(v.get('rmse'))} | {fmt(v.get('mae'))} | {fmt(v.get('mape'))} | {fmt(v.get('inference_time_sec'))}")
    if ('gnn' not in summary) and gnn:
        lines.append(f"gnn | {fmt(gnn.get('rmse'))} | {fmt(gnn.get('mae'))} | {fmt(gnn.get('mape'))} | -")

    lines.append("")
    lines.append("### MLP 特征消融结果表")
    if mlp_variants:
        lines.append("特征 | RMSE | MAE | MAPE | 训练时长(s) | epochs | hidden_dim | lr")
        lines.append("--- | ---: | ---: | ---: | ---: | ---: | ---: | ---:")
        for mv in mlp_variants:
            lines.append(
                f"{mv.get('name')} | {fmt(mv.get('rmse'))} | {fmt(mv.get('mae'))} | {fmt(mv.get('mape'))} | {fmt(mv.get('train_time_sec'))} | {fmt(mv.get('epochs'))} | {fmt(mv.get('hidden_dim'))} | {fmt(mv.get('lr'))}"
            )
    else:
        lines.append("（未找到多特征 MLP 指标文件）")

    lines.append("")
    lines.append("### GNN 训练与推理信息")
    if gnn:
        lines.append("项 | 数值")
        lines.append("--- | ---:")
        lines.append(f"RMSE | {fmt(gnn.get('rmse'))}")
        lines.append(f"MAE | {fmt(gnn.get('mae'))}")
        lines.append(f"MAPE | {fmt(gnn.get('mape'))}")
        lines.append(f"训练时长(s) | {fmt(gnn.get('train_time_sec'))}")
        lines.append(f"推理耗时(s) | {fmt(gnn_inf_time)}")
        lines.append(f"hidden_dim | {fmt(gnn_hidden_dim)}")
        lines.append(f"num_layers | {fmt(gnn_num_layers)}")

    # 图像嵌入（Markdown）
    lines.append("")
    lines.append("### 推理耗时对比")
    if os.path.exists(os.path.join(artifacts, fig_inference)):
        lines.append(f"![Inference Time]({fig_inference})")
    else:
        lines.append("（未找到推理耗时图）")

    lines.append("")
    lines.append("### MLP 指标对比与散点图")
    if os.path.exists(os.path.join(artifacts, fig_mlp_metrics)):
        lines.append(f"![MLP Metrics]({fig_mlp_metrics})")
    if os.path.exists(os.path.join(artifacts, fig_mlp_scatter)):
        lines.append(f"![MLP Scatter]({fig_mlp_scatter})")
    if os.path.exists(os.path.join(artifacts, fig_mlp_error)):
        lines.append(f"![MLP Error]({fig_mlp_error})")

    lines.append("")
    lines.append("### 分距离段误差分箱分析")
    if os.path.exists(os.path.join(artifacts, fig_error_bins)):
        lines.append(f"![Error Bins]({fig_error_bins})")
    else:
        lines.append("（未找到误差分箱图）")

    # 理论与启发式分析
    lines.append("")
    lines.append("## 理论与启发式分析")
    lines.append("- A* 启发式使用地理直线距离（Haversine），在边权为几何距离且满足三角不等式的道路网络上是可采纳且一致的：不会高估真实最短距离，且满足 `h(u) ≤ w(u,v)+h(v)`，因此 A* 保证找到最优路径并能显著减少扩展节点数。")
    lines.append("- 相比 Dijkstra，A* 在具有空间嵌入的图（如道路网络）能以常数因子加速，结合分箱误差图与扩展计数对比，可见在不同距离段均能缩小搜索范围。")

    # 扩展计数对比
    lines.append("")
    lines.append("### 扩展节点与前沿大小对比（A* vs Dijkstra）")
    exp_stats_path = os.path.join(artifacts, 'expansion_stats.json')
    if os.path.exists(exp_stats_path):
        try:
            est = load_json(exp_stats_path) or {}
            agg = est.get('aggregate', {})
            lines.append("项 | 数值")
            lines.append("--- | ---:")
            lines.append(f"样本对数 | {fmt(agg.get('pairs'))}")
            lines.append(f"Dijkstra平均扩展节点 | {fmt(agg.get('dijkstra_expanded_avg'))}")
            lines.append(f"A*平均扩展节点 | {fmt(agg.get('a_star_expanded_avg'))}")
            lines.append(f"Dijkstra平均前沿大小 | {fmt(agg.get('dijkstra_frontier_avg'))}")
            lines.append(f"A*平均前沿大小 | {fmt(agg.get('a_star_frontier_avg'))}")
            lines.append(f"扩展速度提升倍数 | {fmt(agg.get('expansion_speedup'))}")
        except Exception:
            pass
    if os.path.exists(os.path.join(artifacts, fig_expansion)):
        lines.append(f"![Expansion Bars]({fig_expansion})")

    # 讨论与结论
    lines.append("")
    lines.append("## 讨论")
    lines.append("- 速度与精度：GNN 在毫秒级推理上有明显优势；现有配置下误差较高，需调参与特征增强。")
    lines.append("- 调参方向：增加训练轮数/隐藏维度、使用 HuberLoss、log1p 目标、增加 Landmark、改进特征。")
    lines.append("- 工程鲁棒性：对 `NaN/Inf` 特征与不可达样本做清洗，避免训练与评估阶段的数值异常。")

    lines.append("")
    lines.append("## 结论")
    lines.append("在真实道路网络上，GNN 近似最短路径距离的速度优势显著，但需要进一步优化以达到实用精度；Dijkstra/A* 仍是精确参考与上限基线。")

    lines.append("")
    lines.append("## 参考文献")
    lines.append("[1] Hamilton et al., Inductive Representation Learning on Large Graphs (GraphSAGE).")
    lines.append("[2] Dijkstra, A note on two problems in connexion with graphs, 1959.")
    lines.append("[3] Hart et al., A Formal Basis for the Heuristic Determination of Minimum Cost Paths, 1968.")

    return "\n".join(lines)

def render_docx(data_dir: str, out_path: str):
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        print("[report] python-docx 未安装，跳过 Word 生成。")
        return False

    artifacts = os.path.join(data_dir, 'artifacts')
    eval_path = os.path.join(artifacts, 'evaluation_summary.json')
    summary = load_json(eval_path) or {}
    gnn_metrics_path = os.path.join(artifacts, 'gnn_metrics.json')
    gnn = load_json(gnn_metrics_path) or {}
    mlp_variant_files = []
    if os.path.isdir(artifacts):
        for f in os.listdir(artifacts):
            if f.startswith('mlp_metrics_') and f.endswith('.json'):
                mlp_variant_files.append(os.path.join(artifacts, f))
    mlp_variants = []
    for p in sorted(mlp_variant_files):
        m = load_json(p) or {}
        name = os.path.basename(p).replace('mlp_metrics_', '').replace('.json', '')
        mlp_variants.append({
            'name': name,
            'rmse': m.get('rmse'),
            'mae': m.get('mae'),
            'mape': m.get('mape'),
            'train_time_sec': m.get('train_time_sec'),
            'epochs': m.get('epochs'),
            'hidden_dim': m.get('hidden_dim'),
            'lr': m.get('lr'),
        })
    gnn_inf_time = None
    if 'gnn' in summary and isinstance(summary['gnn'], dict):
        gnn_inf_time = summary['gnn'].get('inference_time_sec')

    doc = Document()
    doc.add_heading('基于图神经网络的最短路径距离近似研究', level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('摘要', level=1)
    doc.add_paragraph('本文在真实道路图上比较 Dijkstra、A* 与 GNN（GraphSAGE）在最短路径距离预测上的性能与效率。Dijkstra 与 A* 精度近乎无误差但耗时较高；GNN 推理速度快但当前配置下误差较大，需进一步优化。')

    doc.add_heading('方法', level=1)
    doc.add_paragraph('基线：Dijkstra、A*；模型：GraphSAGE 节点编码 + MLP 距离回归；特征：经纬度、度数、PageRank、路口标识、Landmark 距离；评估：RMSE/MAE/MAPE 与推理耗时。')

    doc.add_heading('实验与结果', level=1)
    # 指标表
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '算法'
    hdr_cells[1].text = 'RMSE'
    hdr_cells[2].text = 'MAE'
    hdr_cells[3].text = 'MAPE'
    hdr_cells[4].text = '推理耗时(s)'
    for key in ['dijkstra', 'a_star', 'mlp', 'gnn']:
        if key in summary:
            v = summary[key]
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = fmt(v.get('rmse'))
            row_cells[2].text = fmt(v.get('mae'))
            row_cells[3].text = fmt(v.get('mape'))
            row_cells[4].text = fmt(v.get('inference_time_sec'))

    if mlp_variants:
        doc.add_heading('MLP 特征消融结果', level=2)
        t2 = doc.add_table(rows=1, cols=8)
        h2 = t2.rows[0].cells
        h2[0].text = '特征'
        h2[1].text = 'RMSE'
        h2[2].text = 'MAE'
        h2[3].text = 'MAPE'
        h2[4].text = '训练时长(s)'
        h2[5].text = 'epochs'
        h2[6].text = 'hidden_dim'
        h2[7].text = 'lr'
        for mv in mlp_variants:
            r = t2.add_row().cells
            r[0].text = str(mv.get('name'))
            r[1].text = fmt(mv.get('rmse'))
            r[2].text = fmt(mv.get('mae'))
            r[3].text = fmt(mv.get('mape'))
            r[4].text = fmt(mv.get('train_time_sec'))
            r[5].text = fmt(mv.get('epochs'))
            r[6].text = fmt(mv.get('hidden_dim'))
            r[7].text = fmt(mv.get('lr'))

    if gnn:
        doc.add_heading('GNN 训练与推理信息', level=2)
        t3 = doc.add_table(rows=1, cols=5)
        h3 = t3.rows[0].cells
        h3[0].text = 'RMSE'
        h3[1].text = 'MAE'
        h3[2].text = 'MAPE'
        h3[3].text = '训练时长(s)'
        h3[4].text = '推理耗时(s)'
        r3 = t3.add_row().cells
        r3[0].text = fmt(gnn.get('rmse'))
        r3[1].text = fmt(gnn.get('mae'))
        r3[2].text = fmt(gnn.get('mape'))
        r3[3].text = fmt(gnn.get('train_time_sec'))
        r3[4].text = fmt(gnn_inf_time)

    # 插入图像（若存在）
    fig_paths = [
        os.path.join(artifacts, 'figures', 'inference_times.png'),
        os.path.join(artifacts, 'figures', 'mlp_metrics.png'),
        os.path.join(artifacts, 'figures', 'mlp_scatter_coords_diff.png'),
        os.path.join(artifacts, 'figures', 'mlp_error_coords_diff.png'),
        os.path.join(artifacts, 'figures', 'error_bins.png'),
        os.path.join(artifacts, 'figures', 'expansion_bars.png'),
    ]
    for p in fig_paths:
        if os.path.exists(p):
            doc.add_picture(p, width=Inches(6))

    doc.add_heading('讨论与结论', level=1)
    doc.add_paragraph('GNN 在毫秒级推理上的优势明显，适合实时近似场景；在当前设置下误差较高，建议通过增加训练轮数与容量、损失函数鲁棒化、目标变换与特征增强等方式进一步提升精度。')

    doc.save(out_path)
    return True

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--data_dir', required=True)
    ap.add_argument('--format', nargs='*', default=['md'], choices=['md', 'docx'])
    args = ap.parse_args()

    artifacts = os.path.join(args.data_dir, 'artifacts')
    os.makedirs(artifacts, exist_ok=True)

    if 'md' in args.format:
        md = build_markdown(args.data_dir)
        md_path = os.path.join(artifacts, 'report.md')
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md)
        print(f"[report] Markdown 已生成: {md_path}")

    if 'docx' in args.format:
        docx_path = os.path.join(artifacts, 'report.docx')
        ok = render_docx(args.data_dir, docx_path)
        if ok:
            print(f"[report] Word 已生成: {docx_path}")

if __name__ == '__main__':
    main()

